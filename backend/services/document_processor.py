# backend/services/document_processor.py

"""
Document Processor Service

Handles extraction of text and metadata from:
- PDF files  (using PyMuPDF)
- DOCX files (using python-docx)
- TXT files  (plain reading)
- CSV files  (using pandas)
- MD files   (plain reading)

Industry practice: One class per responsibility.
This class ONLY processes documents — nothing else.
"""

import os
import re
import fitz          # PyMuPDF — fastest PDF library
import pandas as pd
from pathlib import Path
from docx import Document as DocxDocument
from dataclasses import dataclass, field
from typing import List, Optional
from utils.logger import get_logger

logger = get_logger(__name__)


# ─────────────────────────────────────────────────────────────
# DATA CLASSES — Structured containers for our data
# ─────────────────────────────────────────────────────────────

@dataclass
class PageContent:
    """
    Represents one page of a document.
    
    Why dataclass?
    Clean, readable, no boilerplate __init__ needed.
    """
    page_number: int
    text: str
    char_count: int = 0

    def __post_init__(self):
        # Auto-calculate char count after creation
        self.char_count = len(self.text)


@dataclass
class ProcessedDocument:
    """
    Represents a fully processed document ready for chunking.
    
    This is the standard output of DocumentProcessor.
    Every file type produces this same structure.
    """
    filename: str
    file_type: str
    file_path: str
    total_pages: int
    pages: List[PageContent]
    full_text: str = ""
    total_chars: int = 0
    metadata: dict = field(default_factory=dict)

    def __post_init__(self):
        # Combine all pages into full text
        if not self.full_text:
            self.full_text = "\n\n".join(
                f"[Page {p.page_number}]\n{p.text}"
                for p in self.pages
                if p.text.strip()
            )
        self.total_chars = len(self.full_text)


# ─────────────────────────────────────────────────────────────
# MAIN PROCESSOR CLASS
# ─────────────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Routes files to the correct parser based on extension.
    
    Usage:
        processor = DocumentProcessor()
        doc = processor.process("path/to/file.pdf")
        print(doc.full_text)
        print(doc.total_pages)
    """

    # Supported file types
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".md"}

    def __init__(self):
        logger.info("DocumentProcessor initialized")

    def process(self, file_path: str) -> ProcessedDocument:
        """
        Main entry point. Detects file type and routes to parser.

        Args:
            file_path: Full path to the uploaded file

        Returns:
            ProcessedDocument with text and metadata

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        # Check file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Check file type is supported
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        logger.info(f"Processing file: {path.name} ({extension})")

        # Route to correct parser
        if extension == ".pdf":
            return self._process_pdf(path)
        elif extension == ".docx":
            return self._process_docx(path)
        elif extension in {".txt", ".md"}:
            return self._process_text(path)
        elif extension == ".csv":
            return self._process_csv(path)

    # ─────────────────────────────────────────────────────────
    # PDF PARSER
    # ─────────────────────────────────────────────────────────

    def _process_pdf(self, path: Path) -> ProcessedDocument:
        """
        Extracts text from PDF page by page using PyMuPDF.

        Why PyMuPDF (fitz)?
        - Fastest PDF library in Python
        - Handles scanned PDFs with embedded text
        - Preserves page numbers (critical for citations)
        - Used by companies like Notion and Dropbox
        """
        pages = []

        try:
            # Open the PDF
            pdf_document = fitz.open(str(path))
            total_pages = len(pdf_document)

            logger.info(f"PDF has {total_pages} pages")

            for page_num in range(total_pages):
                # Get page object
                page = pdf_document[page_num]

                # Extract text — "text" mode preserves reading order
                raw_text = page.get_text("text")

                # Clean the extracted text
                cleaned_text = self._clean_text(raw_text)

                # Only add page if it has content
                if cleaned_text.strip():
                    pages.append(PageContent(
                        page_number=page_num + 1,  # Human readable (1-indexed)
                        text=cleaned_text
                    ))

            pdf_document.close()

        except Exception as e:
            logger.error(f"Error processing PDF {path.name}: {e}")
            raise

        return ProcessedDocument(
            filename=path.name,
            file_type="pdf",
            file_path=str(path),
            total_pages=len(pages),
            pages=pages,
            metadata={
                "original_pages": total_pages,
                "pages_with_content": len(pages)
            }
        )

    # ─────────────────────────────────────────────────────────
    # DOCX PARSER
    # ─────────────────────────────────────────────────────────

    def _process_docx(self, path: Path) -> ProcessedDocument:
        """
        Extracts text from Word documents paragraph by paragraph.

        DOCX files don't have pages like PDFs.
        We group every 10 paragraphs into a "virtual page"
        so we can still provide citations.
        """
        pages = []

        try:
            doc = DocxDocument(str(path))

            # Extract all non-empty paragraphs
            paragraphs = [
                p.text.strip()
                for p in doc.paragraphs
                if p.text.strip()
            ]

            logger.info(f"DOCX has {len(paragraphs)} paragraphs")

            # Group paragraphs into virtual pages (10 paragraphs each)
            chunk_size = 10
            for i in range(0, len(paragraphs), chunk_size):
                page_paragraphs = paragraphs[i:i + chunk_size]
                page_text = "\n".join(page_paragraphs)
                cleaned_text = self._clean_text(page_text)

                if cleaned_text.strip():
                    pages.append(PageContent(
                        page_number=(i // chunk_size) + 1,
                        text=cleaned_text
                    ))

        except Exception as e:
            logger.error(f"Error processing DOCX {path.name}: {e}")
            raise

        return ProcessedDocument(
            filename=path.name,
            file_type="docx",
            file_path=str(path),
            total_pages=len(pages),
            pages=pages,
            metadata={"total_paragraphs": len(paragraphs) if 'paragraphs' in locals() else 0}
        )

    # ─────────────────────────────────────────────────────────
    # TXT / MD PARSER
    # ─────────────────────────────────────────────────────────

    def _process_text(self, path: Path) -> ProcessedDocument:
        """
        Reads plain text and markdown files.
        Splits into virtual pages of 3000 characters each.
        """
        pages = []

        try:
            # Read with UTF-8, fallback to latin-1
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                content = path.read_text(encoding="latin-1")

            cleaned = self._clean_text(content)
            logger.info(f"Text file: {len(cleaned)} characters")

            # Split into virtual pages of 3000 chars
            page_size = 3000
            for i in range(0, len(cleaned), page_size):
                chunk = cleaned[i:i + page_size]
                if chunk.strip():
                    pages.append(PageContent(
                        page_number=(i // page_size) + 1,
                        text=chunk
                    ))

        except Exception as e:
            logger.error(f"Error processing text file {path.name}: {e}")
            raise

        return ProcessedDocument(
            filename=path.name,
            file_type=path.suffix.lower().replace(".", ""),
            file_path=str(path),
            total_pages=len(pages),
            pages=pages,
            metadata={"total_characters": len(cleaned) if 'cleaned' in locals() else 0}
        )

    # ─────────────────────────────────────────────────────────
    # CSV PARSER
    # ─────────────────────────────────────────────────────────

    def _process_csv(self, path: Path) -> ProcessedDocument:
        """
        Reads CSV files using pandas.
        Converts each row to readable text format.

        Example:
        Name, Age, City
        John, 25, NYC
        →
        "Row 1: Name=John | Age=25 | City=NYC"
        """
        pages = []

        try:
            df = pd.read_csv(str(path))
            logger.info(f"CSV: {len(df)} rows, {len(df.columns)} columns")

            # Convert dataframe info to text
            summary = (
                f"CSV File: {path.name}\n"
                f"Rows: {len(df)}\n"
                f"Columns: {', '.join(df.columns.tolist())}\n\n"
            )

            # Convert rows to readable text (50 rows per page)
            rows_per_page = 50
            for i in range(0, len(df), rows_per_page):
                chunk_df = df.iloc[i:i + rows_per_page]
                rows_text = []

                for idx, row in chunk_df.iterrows():
                    row_text = " | ".join(
                        f"{col}={val}"
                        for col, val in row.items()
                        if pd.notna(val)
                    )
                    rows_text.append(f"Row {idx + 1}: {row_text}")

                page_text = summary + "\n".join(rows_text)

                pages.append(PageContent(
                    page_number=(i // rows_per_page) + 1,
                    text=page_text
                ))

        except Exception as e:
            logger.error(f"Error processing CSV {path.name}: {e}")
            raise

        return ProcessedDocument(
            filename=path.name,
            file_type="csv",
            file_path=str(path),
            total_pages=len(pages),
            pages=pages,
            metadata={
                "rows": len(df) if 'df' in locals() else 0,
                "columns": df.columns.tolist() if 'df' in locals() else []
            }
        )

    # ─────────────────────────────────────────────────────────
    # TEXT CLEANER
    # ─────────────────────────────────────────────────────────

    def _clean_text(self, text: str) -> str:
        """
        Cleans raw extracted text.

        Why clean?
        PDFs often have:
        - Extra whitespace and blank lines
        - Weird unicode characters
        - Headers/footers repeated on every page
        - Hyphenated words split across lines

        Cleaning improves embedding quality significantly.
        """
        if not text:
            return ""

        # Fix hyphenated line breaks (re-join split words)
        # "impor-\ntant" → "important"
        text = re.sub(r"-\n", "", text)

        # Replace multiple spaces with single space
        text = re.sub(r" +", " ", text)

        # Replace 3+ newlines with 2 newlines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove weird unicode characters but keep common ones
        text = re.sub(r"[^\x00-\x7F\u00C0-\u024F\u1E00-\u1EFF]", " ", text)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text